import fitz  # PyMuPDF
from docx import Document
import re

def clean_text(text):
    """Clean text by removing null bytes and control characters."""
    text = text.replace('\x00', '')  # Remove null bytes
    text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', text)  # Remove control characters
    return text

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF and handle complex layouts."""
    pdf_document = fitz.open(pdf_path)
    full_text = ""

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        # Extract text using different methods for better accuracy
        text = page.get_text("text") or ""
        text = clean_text(text)
        full_text += f"Page {page_num + 1}\n{text}\n\n"  # Add page breaks

    return full_text

def pdf_to_word(pdf_path, word_path):
    """Convert PDF to Word document with improved accuracy."""
    text = extract_text_from_pdf(pdf_path)
    
    # Create a Word document
    doc = Document()

    # Add extracted text to the document
    if text.strip():
        doc.add_paragraph(text)
    else:
        doc.add_paragraph("No text found in the PDF.")

    # Save the Word document
    doc.save(word_path)
    print(f"Conversion successful: {word_path}")
